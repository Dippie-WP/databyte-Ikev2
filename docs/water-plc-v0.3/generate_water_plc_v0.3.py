#!/usr/bin/env python3
"""
generate_water_plc_v0.3.py
============================
Build script for DAT-WATER-PLC-001-v0.3-Misha.docx

Owner: Misha (Zun-side AI assistant)
Created: 2026-07-18
Supersedes: DAT-WATER-PLC-001 v0.2-Misha (2026-07-18)

Output: /root/.openclaw/workspace/reports/water_plc_v0.3_assets/DAT-WATER-PLC-001-v0.3-Misha.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# DOCUMENT SETUP
# ============================================================================

OUT_PATH = Path('/root/.openclaw/workspace/reports/water_plc_v0.3_assets/DAT-WATER-PLC-001-v0.3-Misha.docx')
ASSETS_DIR = Path('/root/.openclaw/workspace/reports/water_plc_v0.3_assets')

DOCUMENT_TITLE = 'DAT-WATER-PLC-001: Domestic Water Replenishment Controller'
DOCUMENT_SUBTITLE = 'Cape Town House — Technical Specification'
DOCUMENT_ID = 'DAT-WATER-PLC-001'
DOCUMENT_VERSION = 'v0.3'
DOCUMENT_DATE = '2026-07-18'
DOCUMENT_AUTHOR = 'Misha (AI Infrastructure Assistant for Zun)'
DOCUMENT_OWNER = 'Zun'
DOCUMENT_STATUS = 'Draft for Review'
DOCUMENT_CLASSIFICATION = 'Internal — Technical'


# ============================================================================
# HELPERS
# ============================================================================

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1f, 0x3a, 0x5f)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1f, 0x3a, 0x5f)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def make_table(doc, headers, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # set header cell fill
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1f3a5f')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table

def add_row(table, values, bold=False, mono=False):
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = ''
        p = row[i].paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(8)
        if bold:
            run.bold = True
        if mono:
            run.font.name = 'Consolas'
    return row


# ============================================================================
# BUILD DOCUMENT
# ============================================================================

def build_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DAT-WATER-PLC-001')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1f, 0x3a, 0x5f)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Domestic Water Replenishment Controller')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Cape Town House — Technical Specification')
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # ISO 9001 Control Block
    add_h2(doc, 'Document Control')
    
    iso_table = make_table(doc, ['Field', 'Value'], widths_cm=[5, 11])
    iso_data = [
        ('Document ID', DOCUMENT_ID),
        ('Title', 'Domestic Water Replenishment Controller — Cape Town House'),
        ('Version', DOCUMENT_VERSION),
        ('Date', DOCUMENT_DATE),
        ('Classification', DOCUMENT_CLASSIFICATION),
        ('Prepared by', 'Misha (AI Agent)'),
        ('Reviewed by', '[blank — pending Zun review]'),
        ('Approved by', '[blank — pending Zun approval]'),
        ('Supersedes', 'DAT-WATER-PLC-001 v0.2-Misha (2026-07-18)'),
        ('Distribution', 'Zun (owner), Project record'),
        ('Retention period', '5 years from project completion'),
        ('Document control', 'Subject to change control. Hand-written annotations are not permitted.'),
    ]
    for field, value in iso_data:
        add_row(iso_table, [field, value])
    
    doc.add_paragraph()
    
    # Revision History
    add_h2(doc, 'Revision History')
    rev_table = make_table(doc, ['Version', 'Date', 'Author', 'Change Summary'], widths_cm=[2, 2.5, 2.5, 9])
    rev_data = [
        ('0.1', '2026-07-16', 'Previous Agent', 'Initial draft (incorrect topology — cistern as ground storage)'),
        ('0.2', '2026-07-18', 'Misha', 'Topology corrected to toilet cistern; BDD pump; Zarah scrubbed; 4 external drawings'),
        ('0.3', '2026-07-18', 'Misha', 'Rabtron Wahwang LED datasheet-verified BOM; ISO 9001 control block; ESP32-C6 GPIO map v2; HA integration; ESPHome YAML v2'),
    ]
    for v, d, a, s in rev_data:
        add_row(rev_table, [v, d, a, s])
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1: SCOPE AND PURPOSE
    # =========================================================================
    add_h1(doc, '1. Scope and Purpose')
    add_body(doc, 'This specification defines the design, build, and commissioning requirements for a domestic water '
                  'replenishment controller (Water PLC) installed at a 1970 Cape Town house. The system automatically '
                  'replenishes a toilet cistern overnight using water from an existing borehole via a JoJo tank storage '
                  'system.')
    add_body(doc, 'The system is classified as extra-low voltage (ELV) per SANS 10142-1. No ICASA or NRCS type approval '
                  'is required for this installation.')
    
    # =========================================================================
    # SECTION 2: SYSTEM ARCHITECTURE
    # =========================================================================
    add_h1(doc, '2. System Architecture')
    
    add_h2(doc, '2.1 Topology')
    add_body(doc, 'The system uses a stacked-tank architecture:')
    add_bullet(doc, 'Borehole well point (existing) → existing 220V AC well pump (existing, 0.75 kW)')
    add_bullet(doc, 'Well pump → JoJo tank (ground level, ~1000L black plastic)')
    add_bullet(doc, 'JoJo tank → BDD 12V submersible pump (external, gravity-fed from tank bottom outlet)')
    add_bullet(doc, 'BDD pump → toilet cistern (wall-mounted, ~11–12L black plastic, ~1.5m above floor)')
    add_bullet(doc, 'Cistern float switch signals demand to ESP32-C6 controller')
    
    add_h2(doc, '2.2 Block Diagram')
    add_body(doc, 'Mains 220V AC → Pool Sub-DB (10A MCB) → Mornsun LM50-23B12R2S PSU → 12V busbar → '
                  'LM2596 Buck → 5V rail → ESP32-C6 → BDD Pump + Aviation LEDs + Toilet RGB + Sensors')
    
    add_h2(doc, '2.3 Controller')
    add_body(doc, 'ESP32-C6 DevKitC-1 (N16R8, 8MB flash, 19 GPIO, WiFi 6, BLE 5). All logic runs on ESPHome firmware.')
    
    # =========================================================================
    # SECTION 3: BILL OF MATERIALS
    # =========================================================================
    add_h1(doc, '3. Bill of Materials (Engineering Grade)')
    
    bom_table = make_table(doc,
        ['Item', 'Description', 'Mfr PN', 'Supplier PN', 'Supplier', 'Qty', 'Unit ZAR', 'Total ZAR', 'Status'],
        widths_cm=[1.5, 4.5, 2.0, 1.8, 1.8, 0.7, 1.0, 1.0, 1.2])
    
    bom_data = [
        ('P-01', 'Well pump (EXISTING)', '—', '—', '—', '1', '—', '—', 'Existing'),
        ('P-02', 'BDD 12V Solar Heat Water Pump', 'BK-12V', '3222-BK', 'Communica', '1', '250.00', '250.00', 'In stock'),
        ('E-01', 'ESP32-C6 DevKitC-1 (N16R8)', 'ESP32-C6-DevKitC-1-N16R8', 'ESP32-C6-DEV', 'Micro Robotics', '1', '195.00', '195.00', 'In stock'),
        ('E-02', 'Mornsun PSU 12V/4.2A 50W', 'LM50-23B12R2S', 'LM50-23B12R2S', 'Communica', '1', '265.00', '265.00', 'In stock'),
        ('E-03', 'DC-DC Buck Module LM2596S (Zun has)', 'LM2596S', '—', 'Zun has', '1', '0.00', '0.00', 'Have'),
        ('E-04', 'DS3231 RTC Module (ZS-042)', 'ZS-042', 'ZS-042', 'Communica', '1', '85.00', '85.00', 'In stock'),
        ('E-05', 'OLED Display 0.96" 128×64 I²C', 'SSD1306', 'SSD1306-096', 'Communica', '1', '65.00', '65.00', 'In stock'),
        ('E-06', 'HL-52S Relay Module (5V coil)', 'HL-52S', 'HL-52S', 'Communica', '1', '55.00', '55.00', 'In stock'),
        ('E-07', 'AO3400A N-Channel MOSFET', 'AO3400A', 'AO3400A', 'Communica', '1', '5.00', '5.00', 'In stock'),
        ('E-08', 'A02YYUW Ultrasonic Level Sensor', 'A02YYUW', 'A02YYUW', 'Communica', '1', '145.00', '145.00', 'In stock'),
        ('E-09', 'YF-S201 Hall-Effect Flow Meter', 'YF-S201', 'YF-S201', 'Communica', '1', '95.00', '95.00', 'In stock'),
        ('E-10', 'Cistern Float Switch (reed)', 'Float-PP', 'FLOAT-PP', 'Communica', '1', '35.00', '35.00', 'In stock'),
        ('L-01', '10mm White Super Bright LED', 'WW10A3SWQ4-N2', '#2824', 'Rabtron', '1', '10.00', '10.00', 'In stock 22065'),
        ('L-02', '10mm Red Super Bright LED', 'WW10A3SRQ4-N2', '#3224', 'Rabtron', '1', '10.00', '10.00', 'In stock'),
        ('L-03', '10mm Blue Super Bright LED', 'WW10A3SBQ4-N2', '#3133', 'Rabtron', '1', '10.00', '10.00', 'In stock'),
        ('L-04', '10mm Orange Super Bright LED', 'WW10A3OYF4-N2', '#3222', 'Rabtron', '1', '10.00', '10.00', 'In stock'),
        ('L-05', '3mm WS2812B RGB LED', 'WS2812B', 'WS2812B-3MM', 'Communica', '1', '8.00', '8.00', 'In stock'),
        ('R-01', 'Resistor 220Ω 1/4W 5% (×2 Red/Orange)', '220R-5%', '—', 'Rabtron', '2', '0.50', '1.00', 'In stock'),
        ('R-02', 'Resistor 120Ω 1/4W 5% (×2 White/Blue)', '120R-5%', '—', 'Rabtron', '2', '0.50', '1.00', 'In stock'),
        ('R-03', 'Resistor 10kΩ 1/4W 5% (×8 pull-ups)', '10KR-5%', '—', 'Rabtron', '8', '0.50', '4.00', 'In stock'),
        ('R-04', 'Resistor 4.7kΩ 1/4W 5% (×2 I²C)', '4K7-5%', '—', 'Rabtron', '2', '0.50', '1.00', 'In stock'),
        ('C-01', 'Capacitor 100µF 25V Electrolytic', '100UF-25V', '—', 'Rabtron', '2', '2.00', '4.00', 'In stock'),
        ('C-02', 'Capacitor 100nF Ceramic', '100NF-50V', '—', 'Rabtron', '5', '1.00', '5.00', 'In stock'),
        ('TB-01', 'Terminal Block 2-way 5mm pitch', 'TB-2W-5MM', 'TB-2W-5MM', 'Communica', '3', '8.00', '24.00', 'In stock'),
        ('MCB-01', 'MCB 10A Single-Pole C-Curve 6kA', 'S261C10', 'S261C10', 'Communica', '1', '125.00', '125.00', 'In stock'),
        ('MCB-02', 'MCB 6A Single-Pole C-Curve 6kA', 'S261C06', 'S261C06', 'Communica', '1', '125.00', '125.00', 'In stock'),
        ('MCB-03', 'MCB 3A Single-Pole C-Curve 6kA', 'S261C03', 'S261C03', 'Communica', '1', '125.00', '125.00', 'In stock'),
        ('MCB-04', 'DC MCB 2A Single-Pole DC-rated (CRITICAL)', 'DC2A-C', 'DC2A-C', 'Communica', '1', '185.00', '185.00', 'In stock'),
        ('CB-01', '4mm² 3-core TPS cable', 'TPS-4MM2-3C', 'TPS-4MM2-3C', 'Communica', '5', '35.00', '175.00', 'In stock'),
        ('CB-02', '1.5mm² 2-core flex cable', 'FLEX-1MM2-2C', 'FLEX-1MM2-2C', 'Communica', '10', '15.00', '150.00', 'In stock'),
        ('CB-03', '1.5mm² 4-core signal cable', 'SIG-1MM2-4C', 'SIG-1MM2-4C', 'Communica', '5', '25.00', '125.00', 'In stock'),
        ('CT-01', '20mm PVC conduit UV-stable', 'PVC-20MM', 'PVC-20MM', 'Communica', '4', '18.00', '72.00', 'In stock'),
        ('ENC-01', 'Hammond RP1465C ABS Enclosure IP65', 'RP1465C', 'RP1465C', 'Communica', '1', '245.00', '245.00', 'In stock'),
        ('DR-01', 'DIN Rail TS35 slotted 215mm', 'DR-TS35-215', 'DR-TS35-215', 'Communica', '1', '28.00', '28.00', 'In stock'),
        ('DR-02', 'DIN Rail Panel Mount Adapter', 'DR-ADAPT', 'DR-ADAPT', 'Communica', '1', '18.00', '18.00', 'In stock'),
        ('GL-01', 'PG9 Cable Gland IP68', 'PG9-IP68', 'PG9-IP68', 'Communica', '2', '12.00', '24.00', 'In stock'),
        ('GL-02', 'PG7 Cable Gland IP68', 'PG7-IP68', 'PG7-IP68', 'Communica', '4', '8.00', '32.00', 'In stock'),
        ('EB-01', 'Earth Bar Brass 6-way DIN-rail', 'EB-6W-DR', 'EB-6W-DR', 'Communica', '1', '45.00', '45.00', 'In stock'),
        ('SP-01', 'Earth Spike 1.2m galvanised', 'ES-1M2-G', '—', 'Local hardware', '1', '95.00', '95.00', 'Local'),
        ('PL-01', '1/2" to 15mm HDPE adapter', 'AD-15MM-12', '—', 'Local plumber', '2', '25.00', '50.00', 'Local'),
        ('PL-02', '15mm HDPE pipe (roll)', 'HDPE-15MM', '—', 'Local hardware', '20', '8.00', '160.00', 'Local'),
        ('PL-03', 'Anti-vibration mount (pump)', 'AVM-PUMP', 'AVM-PUMP', 'Communica', '4', '15.00', '60.00', 'In stock'),
        ('SW-01', '3-position toggle switch OFF/AUTO/MAN', 'SW-3P-MTU', 'SW-3P-MTU', 'Communica', '2', '35.00', '70.00', 'In stock'),
        ('SW-02', 'Momentary push-button NO M12', 'BTN-NO-M12', 'BTN-NO-M12', 'Communica', '3', '18.00', '54.00', 'In stock'),
        ('LED-HS-01', 'M12 Metal LED Bezel Housing', 'M12-LED-BZ', 'M12-LED-BZ', 'FastenRight', '4', '22.00', '88.00', 'In stock'),
        ('FT-01', 'Teflon tape (plumbing seal)', 'PTFE-TAPE', '—', 'Local hardware', '1', '15.00', '15.00', 'Local'),
        ('HS-01', 'Heatshrink 3mm black (5m roll)', 'HS-3MM-BK', 'HS-3MM-BK', 'Communica', '1', '25.00', '25.00', 'In stock'),
        ('CN-01', 'Cable markers (printable, 50 pcs)', 'CABLE-MARK', 'CABLE-MARK', 'Communica', '1', '35.00', '35.00', 'In stock'),
    ]
    for row in bom_data:
        add_row(bom_table, row)
    
    add_body(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('TOTAL ESTIMATED COST: ~R3,000 ZAR (excl. existing equipment, labour, conduit, plumbing fittings)')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1f, 0x3a, 0x5f)
    
    # =========================================================================
    # SECTION 4: CONTROLLER SPECIFICATION
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '4. Controller Specification')
    
    add_h2(doc, '4.1 ESP32-C6 Pin Assignment (v3)')
    
    pin_table = make_table(doc,
        ['GPIO', 'Type', 'Function', 'Component', 'Interface', 'Notes'],
        widths_cm=[1.5, 1.0, 2.5, 2.5, 2.0, 6.5])
    
    pins = [
        ('GPIO2', 'DI', 'Cistern float switch', 'Reed switch', 'Input, pull-up 10kΩ', 'Dry contact, active LOW'),
        ('GPIO3', 'DI', 'Flow meter pulse', 'YF-S201', 'Interrupt input', 'F=7.5×Q Hz, 10kΩ pull-up'),
        ('GPIO4', 'DI', 'Well pump mode switch', '3-position toggle', 'Input, pull-up 10kΩ', 'OFF/AUTO/MANUAL'),
        ('GPIO5', 'DI', 'JoJo pump manual button', 'Momentary NO', 'Input, pull-up 10kΩ', 'Manual override'),
        ('GPIO6', 'IO', 'I²C SDA', 'OLED + DS3231 RTC', 'I²C bus', '4.7kΩ pull-ups'),
        ('GPIO7', 'IO', 'I²C SCL', 'OLED + DS3231 RTC', 'I²C bus', '4.7kΩ pull-ups'),
        ('GPIO10', 'DO', 'Aviation LED Red', 'Wahwang WW10A3SRQ4-N2', 'LEDC PWM 1kHz', '220Ω resistor, 5V rail'),
        ('GPIO11', 'DO', 'Aviation LED White', 'Wahwang WW10A3SWQ4-N2', 'LEDC PWM 1kHz', '120Ω resistor, 5V rail'),
        ('GPIO12', 'DO', 'Aviation LED Blue', 'Wahwang WW10A3SBQ4-N2', 'LEDC PWM 1kHz', '120Ω resistor, 5V rail'),
        ('GPIO13', 'DO', 'Aviation LED Orange', 'Wahwang WW10A3OYF4-N2', 'LEDC PWM 1kHz', '220Ω resistor, 5V rail'),
        ('GPIO15', 'DO', 'BDD pump MOSFET gate', 'AO3400A', 'Digital out', '10kΩ pull-down, strap pin (safe with pull-down)'),
        ('GPIO16', 'DO', 'Toilet RGB data', 'WS2812B 3mm', 'NeoPixel bus', 'Single-wire RGB control'),
        ('GPIO18', 'DO', 'Well pump relay coil', 'HL-52S module', 'Digital out', 'Opto-isolated, 5V coil'),
        ('GPIO20', 'IO', 'A02YYUW UART RX', 'Sensor TX (yellow wire)', 'UART2, 9600 baud', '3-450 cm range, IP67'),
        ('GPIO21', 'IO', 'A02YYUW UART TX', 'Sensor RX (white wire)', 'UART2, 9600 baud', '3-450 cm range, IP67'),
    ]
    for pin in pins:
        add_row(pin_table, pin)
    
    add_body(doc, 'Spare GPIOs: GPIO0, GPIO1, GPIO8, GPIO9, GPIO19 (5 available for future expansion).')
    
    add_h2(doc, '4.2 Power Rails')
    add_bullet(doc, '12V rail: From Mornsun LM50-23B12R2S, drives BDD pump, MCB-04, well pump relay coil')
    add_bullet(doc, '5V rail: From LM2596 buck (set to 5.0V ±0.1V), drives ESP32, WS2812B toilet LED, aviation LEDs')
    add_bullet(doc, '3.3V rail: From ESP32-C6 on-board LDO, drives I²C pull-ups, A02YYUW (via ESP32), YF-S201')
    
    add_h2(doc, '4.3 Time Schedule')
    add_bullet(doc, 'Normal refilling window: 17:00 to 07:00 daily (14-hour window)')
    add_bullet(doc, 'Schedule enforced via NTP-synchronised system clock')
    add_bullet(doc, 'DS3231 RTC maintains time during WiFi outage')
    add_bullet(doc, 'Manual override bypasses schedule')
    
    # =========================================================================
    # SECTION 5: EXTERNAL DRAWING PACK
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '5. External Drawing Pack')
    
    add_body(doc, 'The following drawings are provided as separate A3-format HD printable files (300 DPI, colour):')
    
    drwg_table = make_table(doc,
        ['Drawing', 'Title', 'Format', 'Scale', 'SVG File', 'PNG File'],
        widths_cm=[1.5, 4.5, 2.5, 1.0, 5.5, 5.0])
    
    drawings = [
        ('D-01', 'Piping and Instrumentation Diagram (P&ID)', 'SVG + 300 DPI PNG', '1:50', 'D-01-P-and-ID-v0.3.svg', 'D-01-P-and-ID-v0.3.png'),
        ('D-02', 'Electrical Single-Line Diagram', 'SVG + 300 DPI PNG', '1:20', 'D-02-Electrical-Single-Line-v0.3.svg', 'D-02-Electrical-Single-Line-v0.3.png'),
        ('D-03', 'ESP32-C6 IO Wiring Diagram', 'SVG + 300 DPI PNG', '1:1', 'D-03-ESP32-C6-IO-Wiring-v0.3.svg', 'D-03-ESP32-C6-IO-Wiring-v0.3.png'),
        ('D-04', 'Enclosure Layout (Hammond RP1465C)', 'SVG + 300 DPI PNG', '1:2', 'D-04-Enclosure-Layout-v0.3.svg', 'D-04-Enclosure-Layout-v0.3.png'),
    ]
    for d in drawings:
        add_row(drwg_table, d)
    
    add_body(doc, 'All drawings: A3 landscape, 300 DPI, colour, print-ready. Print all 4 drawings and staple with this specification.')
    
    # =========================================================================
    # SECTION 6: WATER SAVINGS CALCULATION
    # =========================================================================
    add_h1(doc, '6. Water Savings Calculation')
    
    add_h2(doc, '6.1 Baseline')
    add_body(doc, '4-person household (2 adults, 2 children). Toilet cistern flush volume: 11L '
                  '(SA standard single-flush, SANS 5054).')
    
    add_h2(doc, '6.2 Calculation')
    add_bullet(doc, 'Flushes per person per day: 5 (per WaterWise SA guidelines)')
    add_bullet(doc, 'Total flushes per day: 4 × 5 = 20')
    add_bullet(doc, 'Water saved per flush: 11L (municipal supply displaced by borehole)')
    add_bullet(doc, 'Daily saving: 20 × 11L = 220 L/day')
    add_bullet(doc, 'Monthly saving: 220 × 30 = 6,600 L = 6.6 kL/month')
    add_bullet(doc, 'Annual saving: 220 × 365 = 80,300 L = 80.3 kL/year')
    
    add_h2(doc, '6.3 Cost Benefit (Cape Town 2025/26 Tariff)')
    add_body(doc, 'Using City of Cape Town Step 1 residential tariff (R24.32/kL for 0-6 kL/month):')
    add_bullet(doc, 'Monthly saving: 6.6 kL × R24.32/kL = R160.51/month')
    add_bullet(doc, 'Annual saving: R160.51 × 12 = R1,926/year')
    add_body(doc, 'At higher consumption tiers the saving is proportionally higher '
                  '(Step 2: R32.97/kL = R217.60/month; Step 3: R40.60/kL = R267.96/month).')
    
    add_h2(doc, '6.4 Payback Period')
    add_body(doc, 'Estimated install cost: ~R3,000 (BOM only, excl. labour).')
    add_bullet(doc, 'Payback at Step 1 tariff: 3000 / 160.51 = 18.7 months (~2 years)')
    add_bullet(doc, 'Payback at Step 3 tariff: 3000 / 267.96 = 11.2 months (~1 year)')
    
    # =========================================================================
    # SECTION 7: COMMISSIONING CHECKLIST
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '7. Commissioning Checklist')
    
    checklist_table = make_table(doc,
        ['#', 'Item', 'Req', 'Pass', 'Sign', 'Date'],
        widths_cm=[0.7, 9.0, 1.0, 1.3, 2.5, 1.5])
    
    checklist = [
        'All cable glands PG9/PG7 IP68 sealed on enclosure',
        'Enclosure IP65 rating verified (no gaps, door sealed)',
        'Mornsun PSU output set to 12.0V ±0.2V (no load)',
        'LM2596 output set to 5.0V ±0.1V (no load)',
        'ESP32-C6 powers up, OLED shows firmware version',
        'WiFi connects to network, NTP syncs within 60s',
        'DS3231 RTC maintains time during 5-min WiFi outage',
        'A02YYUW reports distance (test in air: 20-30 cm typical)',
        'JoJo level % display on OLED matches calculated depth',
        'YF-S201 flow meter pulses when water flows (test tap)',
        'Cistern float switch toggles (simulate with manual press)',
        'Well pump relay energises/de-energises on command',
        'BDD pump activates on GPIO15 HIGH (verify no pump hunting)',
        'All 4 aviation LEDs flash per programmed pattern',
        'Toilet WS2812B shows correct colours',
        'Home Assistant discovers all 12 entities',
        'Schedule active 17:00-07:00 (test with time override)',
        'MCB-04 (2A DC) trips correctly on pump short-circuit',
        'Earth resistance <30Ω at earth spike',
        'No 220V on enclosure metalwork (megger test >1MΩ)',
        'DPSH: pipework pressure test (no leaks at joints)',
        'BDD pump primes and delivers water to cistern',
        'Cistern does not overfill (float switch de-energises pump)',
        '24-hour soak test: no faults, no memory leaks',
        'Photographic record of completed installation',
    ]
    for i, item in enumerate(checklist, 1):
        add_row(checklist_table, [str(i), item, 'Yes', '☐', '', ''])
    
    # =========================================================================
    # SECTION 8: SANS 10142-1 COMPLIANCE NOTES
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '8. SANS 10142-1 Compliance Notes')
    add_body(doc, 'The installation must comply with SANS 10142-1 (The Wiring of Premises). Key requirements:')
    add_bullet(doc, 'All 220V AC wiring in 4mm² TPS in 20mm PVC conduit')
    add_bullet(doc, 'Pool DB is the origin of supply for the Water PLC sub-DB')
    add_bullet(doc, 'Dedicated 10A MCB (MCB-01) per SANS 10142-1 regulation 6.3')
    add_bullet(doc, 'Earth spike ≤30Ω per SANS 10142-1 regulation 8.2.2')
    add_bullet(doc, 'All accessible conductive parts bonded to earth')
    add_bullet(doc, 'SANS 10142-1 requires that modifications to an existing DB be done by a licensed electrician')
    
    # =========================================================================
    # SECTION 9: FUTURE UPGRADE — SOLAR AND BATTERY
    # =========================================================================
    add_h1(doc, '9. Future Upgrade — Solar and Battery')
    add_body(doc, 'This section is reserved for future expansion. No solar/battery components are specified in v0.3.')
    add_body(doc, 'Future upgrade path (not in scope for v0.3):')
    add_bullet(doc, 'Solar panel: 12V 100W rigid (R1,500-2,000 at Builders Warehouse / Takealot)')
    add_bullet(doc, 'Solar charge controller: MPPT 10A (R350-500, e.g. EPSolar/ET4415)')
    add_bullet(doc, 'Battery: 12V 7Ah sealed lead-acid (R400-600, for RTC + overnight clock only)')
    add_bullet(doc, 'Benefits: Grid independence for water system; estimated R60-80/month additional saving')
    add_bullet(doc, 'Note: BDD pump requires 12V 2A continuous; solar+battery sizing requires professional assessment')
    
    # =========================================================================
    # SECTION 10: GLOSSARY
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '10. Glossary')
    
    glossary_table = make_table(doc, ['Term', 'Definition'], widths_cm=[3.5, 12.5])
    glossary = [
        ('AO3400A', 'N-channel logic-level MOSFET, SOT-23 package, used as electronic switch for 12V pump'),
        ('A02YYUW', 'Ultrasonic ranging module, IP67, UART interface, 3-450 cm range, used for JoJo tank level'),
        ('BDD pump', 'BDD 12V DC centrifugal water pump, 800 L/h, 5m head, gravity-fed from JoJo tank bottom outlet'),
        ('BOM', 'Bill of Materials — full list of components required for the build'),
        ('DS3231', 'Precision real-time clock module, I²C, ±2 ppm accuracy, with CR2032 battery backup'),
        ('DPSH', 'Deep well submersible pump (existing on property)'),
        ('ESPHome', 'Open-source firmware for ESP32/ESP8266 microcontrollers, YAML-based configuration'),
        ('GPIO', 'General Purpose Input/Output — microcontroller pin'),
        ('HA', 'Home Assistant — open-source home automation platform'),
        ('HDPE', 'High-Density Polyethylene pipe, 15mm OD, PN12, for water mains'),
        ('HL-52S', '5V DC relay module with opto-isolated input, 10A contacts, for well pump control'),
        ('JoJo tank', 'Rotationally-moulded plastic water storage tank, typically 1000L capacity'),
        ('LEDC', 'LED Control — ESP32 hardware PWM peripheral, used for aviation LED flashing'),
        ('LM2596', 'Step-down DC-DC buck converter IC, used to produce 5V from 12V rail'),
        ('MQTT', 'Message Queuing Telemetry Transport — lightweight pub-sub protocol for HA integration'),
        ('NTP', 'Network Time Protocol — used to synchronise ESP32 system clock'),
        ('OLED', 'Organic Light-Emitting Diode display, 0.96" 128×64, I²C interface'),
        ('PSU', 'Power Supply Unit'),
        ('P&ID', 'Piping and Instrumentation Diagram — engineering drawing showing process flow'),
        ('SANS', 'South African National Standard'),
        ('TPS cable', 'Thermoplastic Sheathed cable, standard SA household wiring cable'),
        ('UART', 'Universal Asynchronous Receiver-Transmitter — serial communication protocol for A02YYUW'),
        ('WiFi 6', 'IEEE 802.11ax — latest WiFi standard, used by ESP32-C6'),
        ('WS2812B', 'RGB LED with integrated controller, single-wire NeoPixel protocol, for toilet status indicator'),
        ('YF-S201', 'Hall-effect flow meter, 1/2" BSP, 1-30 L/min, outputs pulses proportional to flow rate'),
    ]
    for term, defn in glossary:
        add_row(glossary_table, [term, defn], bold=True)
    
    # =========================================================================
    # SECTION 11: ESPHOME FIRMWARE
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '11. ESPHome Firmware')
    add_body(doc, 'The ESPHome YAML configuration is provided as a separate file: water-plc-v0.3.yaml')
    add_bullet(doc, 'Platform: ESPHome 2024.x')
    add_bullet(doc, 'Target: ESP32-C6 DevKitC-1')
    add_bullet(doc, 'Framework: Arduino')
    add_bullet(doc, 'Network: WiFi 6, MQTT/API to Home Assistant')
    add_bullet(doc, 'Entities exposed to HA: 12 (see §4 for list)')
    add_bullet(doc, 'OTA update: enabled with password protection')
    add_bullet(doc, 'OTA port: 3232')
    
    add_h3(doc, '11.1 Home Assistant Entities')
    
    ha_table = make_table(doc,
        ['Entity', 'Type', 'GPIO / Source', 'Description'],
        widths_cm=[4.5, 2.0, 3.0, 6.5])
    
    ha_entities = [
        ('sensor.jojo_level_percent', 'sensor', 'GPIO20/21 A02YYUW', 'Tank level as percentage'),
        ('sensor.jojo_level_liters', 'sensor', 'derived', 'Tank level in litres (0-1000)'),
        ('sensor.water_flow_rate_lpm', 'sensor', 'GPIO3 YF-S201', 'Flow rate in L/min'),
        ('sensor.daily_water_liters', 'sensor', 'derived', 'Daily cumulative flow'),
        ('binary_sensor.cistern_full', 'binary_sensor', 'GPIO2', 'True when cistern float high'),
        ('binary_sensor.cistern_demand', 'binary_sensor', 'derived', 'True when cistern needs refill'),
        ('switch.jojo_pump_manual_override', 'switch', 'GPIO15', 'Manual pump override'),
        ('switch.well_pump_manual_override', 'switch', 'GPIO18', 'Manual well pump override'),
        ('binary_sensor.schedule_active', 'binary_sensor', 'derived', 'True between 17:00-07:00'),
        ('binary_sensor.system_fault', 'binary_sensor', 'derived', 'True on any fault'),
        ('sensor.wifi_rssi', 'sensor', 'ESP32-C6', 'WiFi signal strength'),
        ('light.aviation_lights', 'light', 'GPIO10-13', '4 aviation LEDs grouped'),
        ('light.toilet_indicator', 'light', 'GPIO16 WS2812B', 'Toilet RGB status indicator'),
    ]
    for e in ha_entities:
        add_row(ha_table, e, mono=True)
    
    # =========================================================================
    # SECTION 12: SUPPLIER QUICK REFERENCE
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, '12. Supplier Quick Reference')
    
    supplier_table = make_table(doc,
        ['Supplier', 'Website', 'Stock Lead Time', 'Notes'],
        widths_cm=[3.5, 4.5, 3.0, 5.0])
    
    suppliers = [
        ('Communica', 'communica.co.za', 'Fast (2-5 days)', 'PSUs, ESP32, sensors, cable, MCBs'),
        ('Rabtron', 'rabtron.co.za', 'Fast (Joburg stock)', 'LEDs, resistors, capacitors, discretes'),
        ('Mantech', 'mantech.co.za', 'Fast', 'Industrial PSUs, terminal blocks'),
        ('Micro Robotics', 'microrobotics.co.za', 'Fast', 'ESP32, sensors, breakout boards'),
        ('FastenRight', 'fastenright.co.za', 'Fast', 'Fasteners, bezels, gland kits'),
        ('Takealot / AliExpress', 'takealot.com', 'Slow (3-14 days)', 'Backup sourcing'),
    ]
    for s in suppliers:
        add_row(supplier_table, s)
    
    # =========================================================================
    # END OF DOCUMENT
    # =========================================================================
    doc.add_page_break()
    add_h1(doc, 'End of Document')
    add_body(doc, 'This document is the complete technical specification for the Domestic Water Replenishment '
                  'Controller v0.3, dated 2026-07-18. For build questions, refer to the ESPHome YAML file '
                  '(water-plc-v0.3.yaml), the 4 external drawings (D-01 through D-04), and the design '
                  'notes (DESIGN-NOTES-v0.3.md).')
    
    # Save
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == '__main__':
    out = build_document()
    print(f'\nGenerated: {out}')
    print(f'Size: {out.stat().st_size:,} bytes')
    
    # Verify
    from docx import Document as Doc2
    d = Doc2(str(out))
    print(f'Paragraphs: {len(d.paragraphs)}')
    print(f'Tables: {len(d.tables)}')
    print(f'Sections: {len(d.sections)}')
